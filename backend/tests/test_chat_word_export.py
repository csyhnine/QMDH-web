import io
import unittest

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.core.security import hash_password
from app.database import Base, get_db
from app.models import ChatMessage, Conversation, User
from app.routers import auth, chat
from app.services.chat_word_export import build_chat_word_document, default_chat_word_file_name


class ChatWordExportTests(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = create_engine(
            "sqlite://",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        Base.metadata.create_all(bind=self.engine)
        self.SessionLocal = sessionmaker(bind=self.engine)

        with self.SessionLocal() as db:
            user = User(
                name="designer.arch",
                display_name="Designer",
                role="designer",
                password_hash=hash_password("qmdh-arch-2026"),
                is_active=True,
                project_codes=["QMDH-001"],
            )
            db.add(user)
            db.commit()
            db.refresh(user)
            conversation = Conversation(user_id=user.id, title="方案讨论", model_provider_id=None)
            db.add(conversation)
            db.commit()
            db.refresh(conversation)
            db.add(
                ChatMessage(
                    conversation_id=conversation.id,
                    role="assistant",
                    content="## 结论\n建议采用方案 A。",
                )
            )
            db.commit()

        self.app = FastAPI()

        def override_get_db():
            with self.SessionLocal() as db:
                yield db

        self.app.dependency_overrides[get_db] = override_get_db
        self.app.include_router(auth.router)
        self.app.include_router(chat.router)
        self.client = TestClient(self.app)

    def login(self) -> str:
        response = self.client.post(
            "/auth/login",
            json={"username": "designer.arch", "password": "qmdh-arch-2026"},
        )
        self.assertEqual(response.status_code, 200, response.text)
        return response.json()["token"]

    def test_build_chat_word_document_contains_content(self) -> None:
        docx_bytes = build_chat_word_document("## 方案摘要\n- 第一点\n正文内容", title="测试对话")
        document = Document(io.BytesIO(docx_bytes))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("测试对话", text)
        self.assertIn("方案摘要", text)
        self.assertIn("第一点", text)
        self.assertIn("正文内容", text)

    def test_default_chat_word_file_name(self) -> None:
        file_name = default_chat_word_file_name(conversation_title="项目讨论", message_id=12)
        self.assertTrue(file_name.endswith(".docx"))
        self.assertIn("项目讨论", file_name)
        self.assertIn("12", file_name)

    def test_export_chat_message_word_endpoint(self) -> None:
        token = self.login()
        headers = {"Authorization": f"Bearer {token}"}
        conversations = self.client.get("/chat/conversations", headers=headers)
        conversation_id = conversations.json()[0]["id"]

        with self.SessionLocal() as db:
            message = db.scalars(select(ChatMessage).where(ChatMessage.role == "assistant")).one()

        response = self.client.post(
            f"/chat/conversations/{conversation_id}/messages/export-word",
            headers=headers,
            json={"message_id": message.id},
        )
        self.assertEqual(response.status_code, 200, response.text)
        self.assertTrue(
            response.headers["content-type"].startswith(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        )
        self.assertIn("attachment", response.headers.get("content-disposition", "").lower())
        document = Document(io.BytesIO(response.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("方案 A", text)


if __name__ == "__main__":
    unittest.main()
