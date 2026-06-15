"""Extract plain text from chat document attachments."""
from __future__ import annotations

import io
from pathlib import Path

from fastapi import HTTPException

from app.services.chat_attachment_upload import read_chat_attachment_bytes

MAX_EXTRACTED_FILE_CHARS = 12_000
_ALLOWED_TEXT_SUFFIXES = {".txt", ".md", ".json", ".csv"}
_ALLOWED_DOCUMENT_SUFFIXES = _ALLOWED_TEXT_SUFFIXES | {".pdf", ".docx", ".xlsx"}


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=503, detail="Word document support is not installed on the server.") from exc

    document = Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _extract_xlsx_text(content: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise HTTPException(status_code=503, detail="Excel spreadsheet support is not installed on the server.") from exc

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sections: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value).strip() for value in row]
                if any(values):
                    rows.append("\t".join(values))
            if rows:
                sections.append(f"[{sheet.title}]\n" + "\n".join(rows))
    finally:
        workbook.close()
    return "\n\n".join(sections)


def extract_attachment_text(storage_path: str, *, file_name: str = "", mime_type: str = "") -> str:
    del mime_type
    normalized_path = storage_path
    suffix = Path(normalized_path).suffix.lower()
    if suffix not in _ALLOWED_DOCUMENT_SUFFIXES:
        raise HTTPException(status_code=400, detail="Unsupported chat document type.")

    content = read_chat_attachment_bytes(normalized_path, allow_documents=True)
    label = file_name.strip() or Path(normalized_path).name

    if suffix in _ALLOWED_TEXT_SUFFIXES:
        extracted = content.decode("utf-8", errors="replace")
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise HTTPException(status_code=503, detail="PDF support is not installed on the server.") from exc
        reader = PdfReader(io.BytesIO(content))
        pages = [str(page.extract_text() or "").strip() for page in reader.pages]
        extracted = "\n\n".join(part for part in pages if part)
    elif suffix == ".docx":
        extracted = _extract_docx_text(content)
    elif suffix == ".xlsx":
        extracted = _extract_xlsx_text(content)
    else:
        raise HTTPException(status_code=400, detail="Unsupported chat document type.")

    extracted = extracted.strip()
    if not extracted:
        raise HTTPException(status_code=400, detail=f"Could not extract text from attachment: {label}")

    if len(extracted) > MAX_EXTRACTED_FILE_CHARS:
        return f"{extracted[:MAX_EXTRACTED_FILE_CHARS].rstrip()}\n...(内容已截断)"
    return extracted


def format_file_attachment_context(file_name: str, extracted_text: str) -> str:
    return f"附件：{file_name}\n{extracted_text.strip()}"
