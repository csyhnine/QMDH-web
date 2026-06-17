"""Build Word documents from chat assistant replies."""
from __future__ import annotations

import io
import re
from datetime import datetime

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_HEADING_PATTERN = re.compile(r"^(#{1,3})\s+(.+)$")
_BULLET_PATTERN = re.compile(r"^[-*]\s+(.+)$")
_ORDERED_PATTERN = re.compile(r"^\d+\.\s+(.+)$")
_LINK_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_BOLD_PATTERN = re.compile(r"\*\*([^*]+)\*\*")
_ITALIC_PATTERN = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
_INLINE_CODE_PATTERN = re.compile(r"`([^`]+)`")
_EXPORT_FONT_NAME = "Microsoft YaHei"


def _safe_export_file_name(value: str, *, fallback: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in {"-", "_", ".", " "} else "-" for char in value.strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .-_")
    if not cleaned:
        cleaned = fallback
    if not cleaned.lower().endswith(".docx"):
        cleaned = f"{cleaned}.docx"
    return cleaned[:120]


def _strip_inline_markdown(text: str) -> str:
    cleaned = _LINK_PATTERN.sub(r"\1", text)
    cleaned = _BOLD_PATTERN.sub(r"\1", cleaned)
    cleaned = _ITALIC_PATTERN.sub(r"\1", cleaned)
    cleaned = _INLINE_CODE_PATTERN.sub(r"\1", cleaned)
    return cleaned.strip()


def _normalize_export_content(content: str) -> str:
    normalized = content.replace("\r\n", "\n").replace("\r", "\n")
    lines: list[str] = []
    in_code_block = False

    for raw_line in normalized.split("\n"):
        line = raw_line.strip()
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            lines.append(line)
            continue
        if not line:
            lines.append("")
            continue
        lines.append(_strip_inline_markdown(line))

    return "\n".join(lines).strip()


def _set_run_font(run, *, font_name: str = _EXPORT_FONT_NAME) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _configure_document_fonts(document, *, font_name: str = _EXPORT_FONT_NAME) -> None:
    for style_name in (
        "Normal",
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Paragraph",
        "List Bullet",
        "List Number",
    ):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:eastAsia"), font_name)
        rfonts.set(qn("w:cs"), font_name)


def _append_paragraph(document, text: str, *, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    if text:
        run = paragraph.add_run(text)
        _set_run_font(run)
    elif not paragraph.runs:
        run = paragraph.add_run("")
        _set_run_font(run)


def _write_markdownish_lines(document, content: str) -> None:
    for raw_line in content.split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            continue

        heading_match = _HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            _append_paragraph(document, heading_match.group(2).strip(), style=style)
            continue

        bullet_match = _BULLET_PATTERN.match(stripped)
        if bullet_match:
            _append_paragraph(document, bullet_match.group(1).strip(), style="List Bullet")
            continue

        ordered_match = _ORDERED_PATTERN.match(stripped)
        if ordered_match:
            _append_paragraph(document, ordered_match.group(1).strip(), style="List Number")
            continue

        _append_paragraph(document, stripped)


def build_chat_word_document(content: str, *, title: str = "") -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Word export support is not installed on the server.") from exc

    document = Document()
    _configure_document_fonts(document)

    heading = _strip_inline_markdown(title.strip() or "Chat 回复")
    _append_paragraph(document, heading, style="Title")
    _append_paragraph(document, datetime.now().strftime("%Y-%m-%d %H:%M"))
    _append_paragraph(document, "")
    _write_markdownish_lines(document, _normalize_export_content(content))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def default_chat_word_file_name(*, conversation_title: str = "", message_id: int | None = None) -> str:
    stub = conversation_title.strip() or "chat-reply"
    if message_id is not None:
        stub = f"{stub}-{message_id}"
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    return _safe_export_file_name(f"{stub}-{timestamp}", fallback=f"chat-reply-{timestamp}.docx")
